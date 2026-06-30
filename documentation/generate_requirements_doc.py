# -*- coding: utf-8 -*-
"""Generate a comprehensive Word guide explaining what is ACTUALLY required
for every sub-item of all 23 CHN1A data-center permits."""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x5b, 0x4a, 0xb0)
DARK   = RGBColor(0x20, 0x24, 0x40)
GREY   = RGBColor(0x55, 0x55, 0x66)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'),'clear'); sh.set(qn('w:fill'),hexcolor)
    tcPr.append(sh)

doc = Document()
# base style
st = doc.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
for s in ['Heading 1','Heading 2','Heading 3','Title']:
    try: doc.styles[s].font.name='Calibri'
    except: pass

def H(text, level=1, color=ACCENT, size=None):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = color
        if size: r.font.size = Pt(size)
    return p

def para(text, bold=False, italic=False, color=None, size=None, space_after=4):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold=bold; r.italic=italic
    if color: r.font.color.rgb=color
    if size: r.font.size=Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def item(name, explanation):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(name + " — "); r.bold=True
    p.add_run(explanation)
    p.paragraph_format.space_after = Pt(4)

def facts(rows):
    t = doc.add_table(rows=0, cols=2); t.alignment=WD_TABLE_ALIGNMENT.LEFT
    t.style='Light Grid Accent 1'
    for k,v in rows:
        c = t.add_row().cells
        c[0].text=k; c[1].text=v
        for r in c[0].paragraphs[0].runs: r.bold=True
        c[0].width=Inches(1.7); c[1].width=Inches(4.6)
    doc.add_paragraph()

def pitfalls(items):
    para("Common pitfalls & review points", bold=True, color=DARK, space_after=2)
    for x in items:
        p = doc.add_paragraph(style='List Bullet 2'); p.add_run(x)
        p.paragraph_format.space_after = Pt(2)
    doc.add_paragraph()

# ============================ COVER ============================
ttl = doc.add_heading('Thailand Data Center Permitting', level=0)
for r in ttl.runs: r.font.color.rgb=ACCENT
para("What each permit submission ACTUALLY requires — a sub-item-level review guide",
     bold=True, color=DARK, size=13)
para("Project: CHN1A GFL Cut-In Data Center", italic=True, color=GREY)
para("23 permits · 13+ authorities · prepared from the project research package", italic=True, color=GREY)
para("Document 08 — companion to the Master Permitting Guide, EIA & FBL guides, and the 23 permit checklists.",
     italic=True, color=GREY, size=9)
doc.add_paragraph()

# ============================ HOW TO READ ============================
H("How to use this guide", 1)
para("For every permit, this guide lists each document/sub-item the authority asks for and explains "
     "what that item actually is, what it must contain, who prepares or issues it, and the format/"
     "legalisation expected. The aim is to remove the guesswork between a checklist line ('Articles of "
     "Association') and a submission-ready document. Use it alongside the per-permit Excel checklist to "
     "track status, and the Master Permitting Guide for the legal background.")

H("Document standards that apply to almost everything", 2)
para("Several requirements recur across nearly all permits. Get these right once and most submissions go smoothly.")
item("Certified true copies", "Copies of company or government documents must usually be signed 'certified true copy' "
     "by an authorised director (with the company seal) or by the issuing officer. Authorities reject plain photocopies.")
item("Foreign documents — apostille or consular legalisation", "Any document issued abroad (certificate of "
     "incorporation, good-standing certificate, passport copy, criminal-record certificate, parent-company financials) "
     "must be authenticated. For Hague-Convention countries an apostille is enough; for others, legalisation by the "
     "Thai embassy/consulate in the issuing country plus the foreign ministry is required.")
item("Certified Thai translation", "Documents not in Thai generally need a translation by a certified translator, and "
     "for court/registry use the translation itself may need to be notarised. Budget 1–2 weeks for translation + legalisation.")
item("Power of Attorney (POA)", "If an agent (law firm, consultant) submits on the company's behalf, a POA naming that "
     "agent and the specific filing is required, signed by an authorised director, sealed, and — if signed abroad — notarised "
     "and apostilled. A generic POA is often rejected; it should reference the exact application.")
item("Professional seals & licences", "Engineering/architectural drawings must carry the seal, signature and licence "
     "number of a Thai-licensed (Council of Engineers / Architect Council) professional. Foreign-stamped drawings are not "
     "accepted on their own; a local licensed professional must take responsibility.")
item("Validity windows", "Company affidavits, shareholder lists and good-standing certificates are usually only accepted "
     "if issued within the last 3–6 months. Pull fresh copies just before submission.")
doc.add_paragraph()

# ============================ PERMIT DATA ============================
# Each permit: title, facts list, purpose, items[(name,explanation)], pitfalls[]
PERMITS = [
("1. Company Registration",
 [("Authority","Department of Business Development (DBD), Ministry of Commerce"),
  ("Phase","Development"),("Lead time","~30 days"),("Est. cost","~5,000 THB"),
  ("Prerequisite for","Everything — this creates the legal entity")],
 "Establishes the Thai juristic person (usually a private limited company) that will hold every other permit, "
 "the land rights and the BOI promotion. Nothing else can proceed until this exists.",
 [("Completed registration application (Bor.Or.Jor.)","The DBD incorporation form set: memorandum of association, "
   "statutory meeting minutes and the registration application. Filed online via the DBD e-Registration portal or at the DBD office."),
  ("Memorandum & Articles of Association","The company's constitution: objectives (must explicitly cover data-centre / "
   "IT-infrastructure operation), registered capital, share structure and director powers. The objectives clause matters later "
   "for BOI and the Foreign Business License, so draft it broadly enough."),
  ("Registered capital evidence","The amount of capital stated in the MOA. For a foreign-majority/BOI project plan capital "
   "realistic to the investment; capital also drives work-permit quotas (≈2M THB per foreign work permit)."),
  ("Director & shareholder identification","National ID (Thai) or passport (foreign) for every director and shareholder, plus "
   "a shareholder list (Bor.Or.Jor.5) showing names, nationalities and exact shareholding percentages. Foreign IDs need certification."),
  ("Registered office proof","A lease or title deed plus a consent letter and house-registration (Tabien Baan) for the registered "
   "address, with a location map. A virtual address is not accepted for a facility operator."),
  ("Company seal","The official rubber stamp; its impression is filed and then used to certify every later document.")],
 ["Make the objectives clause explicitly include data-centre operation, leasing of IT infrastructure and related services — "
  "amending it later delays BOI and FBL.",
  "Foreign shareholders trigger the Foreign Business Act: if foreigners hold >49%, you need the FBL (Permit 3) or a BOI card.",
  "Keep at least 3 fresh certified copies of the company affidavit (Nangsue Rabrong) — almost every other authority asks for one."]),

("2. BOI Promotion Certificate",
 [("Authority","Thailand Board of Investment (BOI)"),
  ("Phase","Development"),("Lead time","~120 days"),("Est. cost","No application fee"),
  ("Prerequisite","Power Availability letter from ERC/MEA/PEA (now mandatory)")],
 "Grants the investment-promotion privileges (tax holidays, 100% foreign ownership, land-holding and work-permit facilitation) "
 "under the data-centre category 8.2.1.1 / 8.2.1.2. The certificate underpins the project's whole financial case.",
 [("Application form (F PA PP 04, digital-technology activity)","The BOI online application describing the activity, capacity (MW/IT "
   "load), location, capital and employment. Category 8.2.1.1 carries a PUE-efficiency commitment; 8.2.1.2 is the standard tier — "
   "choose deliberately because the obligations differ."),
  ("ERC/MEA/PEA Power Supply Confirmation letter","A letter from the energy regulator/utility confirming grid power of the required "
   "capacity is available to the site. As of 2025–26 this is a hard prerequisite — BOI will not accept the application without it. "
   "This, not the BOI form, is usually the real bottleneck (2–8 weeks)."),
  ("Three-year financial projections / feasibility study","A model showing investment, revenue, cost and viability. For investments "
   "above ~2 bn THB a full feasibility study is expected. Numbers must reconcile with the registered capital and the build programme."),
  ("Technical specification pack","Single-line electrical diagrams, HVAC/cooling design, network topology and — for 8.2.1.1 — PUE "
   "modelling and CFD studies evidencing the efficiency target. These show the project is technically real, not a shell."),
  ("Thailand Benefit Enhancement Plan","The local-value commitments BOI now scrutinises: university MOUs for curriculum, a training "
   "plan (often a 10× workforce multiplier) and SME/local-supply-chain integration. Tax activation is later checked against this plan."),
  ("Shareholder / parent-company evidence","Shareholder list with nationalities, and for corporate investors the parent's audited "
   "financial statements and certificate of incorporation (apostilled). For individual foreign investors, a clean criminal-record "
   "certificate and home-country tax ID (apostilled + translated).")],
 ["The ERC power letter gates the whole application — start it first, in parallel with company registration.",
  "Tax exemption does NOT switch on at approval: CIT relief is verified later against ISO/IEC 27001, the Benefit Plan and (for 8.2.1.1) the PUE result.",
  "From 2026 monitoring moved to quarterly e-Monitoring; two consecutive missed reports can revoke the certificate — set the reporting calendar at day one."]),

("3. Foreign Business License (FBL)",
 [("Authority","Department of Business Development (DBD) — Foreign Business Committee"),
  ("Phase","Development"),("Lead time","~60 days (often longer)"),("Est. cost","~10,000 THB + capital"),
  ("When needed","Foreign ownership >49% and no BOI card covering the activity")],
 "Authorises a foreign-majority company to carry on a restricted-list business in Thailand. If BOI promotion already grants a "
 "Foreign Business Certificate for the activity, a separate FBL may not be needed — confirm overlap before filing both.",
 [("Application form P.13","The FBL application describing the foreign business activity, capacity and how it benefits Thailand "
   "(technology transfer, jobs, local spend). The 'benefit to Thailand' narrative is decisive for the committee."),
  ("Business plan in Thai (5–10 pp)","A detailed plan: services, operations, customers, revenue and management. Must be in Thai and "
   "specific — a generic plan is the most common rejection reason."),
  ("Minimum capital evidence","Proof of at least 2,000,000 THB (3M+ if the activity is on List 2/3) actually remitted and held for "
   "the prescribed period, with a bank certification letter. The money must genuinely be in the account, not just pledged."),
  ("Shareholder register & affidavit","Bor.Or.Jor.5 and the company affidavit (issued within 3–6 months) showing the foreign "
   "shareholding that triggers the FBL."),
  ("Foreign investor documents","Passport copy, certificate of no criminal record and home tax documents — each apostilled and "
   "translated to Thai. For a corporate foreign shareholder, the good-standing certificate and board resolution to invest.")],
 ["Check first whether the BOI card already lets you operate — duplicating with an FBL wastes 2–3 months.",
  "The 'benefit to Thailand' section is graded; tie it to the BOI Benefit Plan for consistency.",
  "Capital must be remitted from abroad and traceable; domestic transfers can be queried."]),

("4. Land Use & Zoning Confirmation",
 [("Authority","Dept. of Public Works & Town/Country Planning / Local Authority (Or.Bor.Tor. / municipality)"),
  ("Phase","Development"),("Lead time","~30 days"),("Est. cost","~2,000 THB"),
  ("Prerequisite for","EIA and the Building Permit")],
 "Confirms the parcel's zoning (under the city/comprehensive plan) permits a data-centre/industrial use with the intended "
 "building scale, setbacks and utilities. A negative zoning finding can stop the project before design begins.",
 [("Land title deed (Chanote / Nor Sor 3 Gor)","A certified copy of the strongest available title. Chanote (full freehold survey) "
   "is preferred; weaker titles can limit financing and BOI land-holding rights. Pull a fresh title search (<3 months)."),
  ("Zoning certificate request","The application asking the planning authority to confirm the land-use colour-zone and what it "
   "allows. Industrial 'purple' or appropriate zones permit data centres; residential/agri zones may not."),
  ("Site plan with boundaries & setbacks","A scaled plan showing the parcel, proposed building footprint, setbacks from boundaries "
   "and roads, and indicative utility routes — enough for the planner to test against zone rules."),
  ("Surrounding land-use context","Notes/photos on neighbours (residences, waterways, protected areas) because buffer and "
   "environmental rules tighten near sensitive uses."),
  ("Lease (if not owner)","A 30-year-plus registered lease or ownership proof — short leases undermine BOI land rights and lender comfort.")],
 ["Verify the parcel is not in a flood-control, military or environmentally-protected overlay that overrides the base zone.",
  "Confirm road access and the legal right-of-way now; landlocked parcels need a registered servitude.",
  "If zoning is marginal, get the planner's written view before committing to design spend."]),

("5. Environmental Impact Assessment (EIA)",
 [("Authority","Office of Natural Resources & Environmental Policy and Planning (ONEP)"),
  ("Phase","Development"),("Lead time","~180 days (the critical-path bottleneck)"),("Est. cost","~200,000–500,000 THB"),
  ("Prerequisite for","Building Permit, water/wastewater, generator, factory licence")],
 "A formal study (by a licensed EIA consultant) of the project's environmental and social impacts and the measures to mitigate "
 "them, approved by ONEP. It is the single longest permit and gates most of the construction phase — start it the moment zoning is clear.",
 [("Screening & Terms of Reference (TOR)","First ONEP confirms the project type and the study scope (TOR). Getting the TOR agreed "
   "early avoids re-work; the TOR defines exactly which baselines and impacts must be studied."),
  ("Baseline environmental data","Measured existing conditions — ambient air quality, noise, water, ecology and socioeconomics — "
   "collected over a defined period. This is real fieldwork/monitoring, not desk estimates, and drives the timeline."),
  ("Impact assessment & modelling","Quantified construction- and operation-phase impacts: generator stack emissions (NOx, PM, SO2) "
   "vs. Thai air standards, noise at the boundary, water demand and wastewater characteristics, traffic and waste."),
  ("Mitigation & Environmental Management Plan","Specific, costed measures (e.g. generator stack height, acoustic enclosures, "
   "wastewater treatment spec) and who is responsible. Vague mitigation is the top revision cause."),
  ("Monitoring plan","The ongoing compliance monitoring (quarterly during construction, at least annually in operation) the company "
   "commits to — this becomes a binding operating obligation."),
  ("Public/community consultation record","Minutes, attendance and photos of consultation meetings, plus written responses to "
   "community concerns. Inadequate consultation gets reports sent back.")],
 ["Engage an ONEP-licensed EIA firm; the report must be authored under their licence.",
  "Baseline monitoring sets the floor on schedule — you cannot compress the measurement window, so begin immediately.",
  "Align the generator and wastewater specs in the EIA with Permits 12 and 14 so the later submissions match what ONEP approved."]),

("6. Building Construction Permit",
 [("Authority","Local Building Control Authority (municipality / Or.Bor.Tor.) under the Building Control Act"),
  ("Phase","Design & Planning"),("Lead time","~60 days"),("Est. cost","~15,000 THB"),
  ("Gates","All construction-phase permits")],
 "Legal authorisation to construct the building to the approved drawings. No site works may start until it is issued; it depends "
 "on zoning and the EIA approval.",
 [("Permit application (Or.1) + land documents","The Building Control application with title deed, zoning confirmation and, where "
   "required, the EIA approval letter attached."),
  ("Architectural drawings","Stamped floor plans, elevations, sections and roof plans at proper scale, with door/window schedules — "
   "showing the building complies with height, setback, coverage and means-of-escape rules."),
  ("Structural drawings & calculations","Foundation and frame design with load and (where relevant) seismic calculations, sealed by a "
   "licensed structural engineer who takes legal responsibility."),
  ("MEP drawings","Mechanical, electrical and plumbing design — HVAC capacity, electrical single-line with load schedule, water/drainage "
   "— consistent with the architectural set."),
  ("Fire & life-safety plan","Escape routes and distances, fire-rated compartments/doors, detection and suppression layout and emergency "
   "lighting — for a data centre these are scrutinised heavily."),
  ("Licensed professional seals","Each drawing set carries the seal, signature and licence number of the responsible Thai-licensed "
   "architect/engineer; this is mandatory, not a formality.")],
 ["Drawings must be internally consistent — architectural, structural and MEP sets that disagree are the main rejection cause.",
  "Confirm whether the EIA approval letter must accompany the application in your locality (it usually must).",
  "Large/high-rise or special buildings may need extra structural review — check the local threshold early."]),

("7. Road Access / Highway Connection",
 [("Authority","Local Authority / Department of Highways or Rural Roads"),
  ("Phase","Design & Planning"),("Lead time","~45 days"),("Est. cost","~8,000 THB"),
  ("Depends on","Building Permit / site layout")],
 "Approval to create or use an access connection from the site onto the public road network, including any works on the road reserve.",
 [("Access application & site/location plan","Request plus plans showing the connection point, internal circulation, parking, loading "
   "and emergency access relative to the public road."),
  ("Traffic impact study","Estimated construction- and operation-phase vehicle volumes, peak flows and the effect on the existing road, "
   "with proposed controls — scaled to the road's capacity."),
  ("Access/road engineering design","Driveway geometry, widths, pavement, sight-lines at the junction and drainage under the access, "
   "sealed by an engineer."),
  ("Land-rights & right-of-way proof","Title/lease and, if the access crosses other land, a registered servitude.")],
 ["Confirm who owns the abutting road (Highways vs. local) — it changes the authority and the standard.",
  "Sight-distance and junction safety are common sticking points; design to the road authority's standard, not the minimum.",
  "Heavy-equipment delivery routes for construction may need separate temporary approval."]),

("8. Grid Connection Approval",
 [("Authority","Provincial Electricity Authority (PEA) or Metropolitan Electricity Authority (MEA)"),
  ("Phase","Design & Planning"),("Lead time","~60 days"),("Est. cost","~50,000–100,000 THB"),
  ("Linked to","Power Availability (ERC) and HV Substation")],
 "The utility's agreement to connect the facility's load to the grid at the required capacity and voltage, defining the point of "
 "connection and the works each party builds.",
 [("Connection application & load schedule","Formal request stating total demand (MW) and profile, broken into IT, cooling, lighting "
   "and auxiliary loads with peak/average figures."),
  ("Single-line electrical diagram","The connection topology from the utility point through the substation/transformers to distribution "
   "— sealed by a licensed electrical engineer."),
  ("Substation/transformer design","Location, capacity, voltage and configuration for the on-site receiving station (cross-references Permit 10)."),
  ("Site map & land control","Showing the proposed substation and cable route, with proof of land rights for that footprint."),
  ("Programme","Construction and energisation dates so the utility can plan its reinforcement works.")],
 ["The ERC Power Availability letter (Permit 9) should already confirm capacity exists — grid connection turns that into a contracted point of supply.",
  "Utility-side reinforcement can have a long lead; lock the connection date against the build programme.",
  "Clarify the cost-sharing and deposit for utility works up front."]),

("9. Power Availability Confirmation (ERC)",
 [("Authority","Energy Regulatory Commission (ERC), via PEA/MEA"),
  ("Phase","Design & Planning"),("Lead time","~30 days"),("Est. cost","No fee"),
  ("Critical","Mandatory PREREQUISITE for the BOI application")],
 "An official confirmation that grid power of the required capacity is available to the site. Since 2025–26 it must be obtained "
 "before the BOI application — making it one of the earliest and most schedule-critical actions.",
 [("Power requirement letter & load forecast","A letter stating the capacity needed (MW) with peak and average load and the connection "
   "timeline."),
  ("Project specification & location","Facility type, size and GPS location so the utility can test capacity against the local network."),
  ("Grid connection plan / single-line","Indicative connection point, transformer and substation arrangement."),
  ("Applicant identity & land proof","Company affidavit and title/lease showing the applicant controls the site."),
  ("Environmental status","Reference to the EIA stage where the utility requires it.")],
 ["Start this FIRST — it gates BOI and can take 2–8 weeks depending on local network headroom.",
  "If local capacity is short, the utility may require reinforcement at the project's cost — surface this early.",
  "Keep the figures consistent with the BOI load numbers and the grid-connection application."]),

("10. High-Voltage Substation Approval",
 [("Authority","PEA / MEA"),
  ("Phase","Design & Planning"),("Lead time","~90 days"),("Est. cost","~100,000–200,000 THB"),
  ("Depends on","Grid Connection Approval")],
 "Technical approval of the on-site HV substation that steps utility voltage down for the facility — design, protection, safety and siting.",
 [("Substation design package","Layout, single-line diagram, transformer ratings/cooling, switchgear and protection equipment, grounding "
   "and cable routing — sealed by a licensed electrical engineer."),
  ("Engineering studies","Load-flow, fault-current and protection-coordination studies, plus equipment thermal analysis, proving the design is safe and selective."),
  ("Safety & security plan","Physical security, fencing, hazard signage, access control and emergency procedures for the HV compound."),
  ("Environmental & siting evidence","Assessment of EMF, noise and visual impact and proof of land control for the substation footprint."),
  ("Building permit (if enclosed)","If the substation sits in a building, its own building approval.")],
 ["Protection-coordination and fault studies are where approvals stall — have them done by an experienced HV engineer.",
  "Coordinate the substation programme with the utility's grid works so energisation isn't held up.",
  "Align ratings with the BOI/grid load figures."]),

("11. Water Supply Connection",
 [("Authority","Metropolitan (MWA) or Provincial (PWA) Waterworks Authority"),
  ("Phase","Construction"),("Lead time","~45 days"),("Est. cost","~20,000–40,000 THB"),
  ("Depends on","Building Permit")],
 "Connection to mains water at the demand the facility needs — significant for a data centre because of evaporative/cooling water use.",
 [("Connection application & demand calculation","Daily and peak water demand split into cooling and domestic use, with seasonal variation."),
  ("Water system drawings","Meter location, pipe sizing and routing from the main, storage tank capacity, backflow prevention and any pre-treatment."),
  ("Land rights & building permit","Title/lease and the building approval."),
  ("Environmental reference","Consistency with the EIA water-use figures.")],
 ["Cooling water demand is large and peaky — size the storage and connection for peak, not average.",
  "Backflow prevention is mandatory for industrial connections; show it explicitly.",
  "If mains capacity is limited, plan on-site storage or an alternative source early."]),

("12. Wastewater / Drainage Connection",
 [("Authority","Local Authority / Waterworks / under the Factory & Public Health rules"),
  ("Phase","Construction"),("Lead time","~45 days"),("Est. cost","~25,000–50,000 THB"),
  ("Depends on","Building Permit, EIA")],
 "Approval to treat and discharge the facility's wastewater to the required effluent standard, and to connect to the public sewer or "
 "receiving water.",
 [("Discharge application & wastewater characterisation","Estimated volume and composition (BOD, COD, TSS, pH, temperature, any "
   "chemical content from cooling/cleaning)."),
  ("Treatment system design","Process (settling, treatment, filtration), capacity, the effluent quality it achieves vs. Thai standards, "
   "and sludge management."),
  ("Drainage layout","Pipe sizing/routing to the discharge point, inspection access and any backup power for treatment."),
  ("Monitoring plan","Effluent testing schedule and reporting committed in the EIA."),
  ("Land rights & building permit","Title/lease and the building approval.")],
 ["Effluent must meet the discharge standard for the receiving environment — design the treatment to that, not a generic spec.",
  "The treatment design must match what the EIA approved; divergence triggers re-review.",
  "Provide backup power for treatment so discharge stays compliant during outages."]),

("13. Fire Protection System Approval",
 [("Authority","Local Building Control / Fire Authority"),
  ("Phase","Construction"),("Lead time","~30 days"),("Est. cost","~80,000–150,000 THB"),
  ("Depends on","Building Permit")],
 "Design approval of the detection and suppression systems protecting the facility — critical for a data centre, where clean-agent "
 "suppression and compartmentation protect IT halls.",
 [("Fire safety plan","Overall strategy: zones, compartmentation, escape strategy and occupancy load."),
  ("Detection design","Smoke/heat detector layout and spacing, manual call points and the control panel — to the relevant standard."),
  ("Suppression design","Agent type (clean agent such as FM-200/IG-541/Novec, or sprinkler where appropriate), nozzle layout, discharge/"
   "concentration calculations and agent storage — sized to the protected volume."),
  ("Emergency lighting & signage","Exit signs and emergency lighting coverage and egress wayfinding."),
  ("Passive protection","Fire-rated walls/doors and compartment boundaries with rated penetrations."),
  ("Licensed fire engineer seal","The design is sealed by a qualified fire-protection engineer.")],
 ["Specify clean-agent suppression for IT halls (water risks equipment); justify the agent and concentration with calculations.",
  "Detection and suppression zoning must match the room compartmentation in the building set.",
  "Coordinate smoke control and stairwell pressurisation with the MEP design."]),

("14. Generator & Fuel Storage Approval",
 [("Authority","Department of Energy Business (DOEB)"),
  ("Phase","Construction"),("Lead time","~60 days"),("Est. cost","~10,000–20,000 THB"),
  ("Depends on","Building Permit, EIA emissions limits")],
 "Approval for the standby generator(s) and associated fuel storage — covering emissions, noise, siting and the fuel-handling safety case.",
 [("Generator specification","Capacity, fuel type, emissions and sound rating, dimensions and the siting (clearances from building and boundary, exhaust routing)."),
  ("Emissions compliance","Stack-emission estimates or test data (NOx, PM10, SO2) compared with Thai air-quality limits and the EIA commitments."),
  ("Noise mitigation","Boundary noise prediction vs. the noise standard, with enclosure/barrier design."),
  ("Fuel storage & secondary containment","Tank capacity, construction and location, with secondary containment sized to at least 110% of "
   "the largest tank, plus spill-response and fire protection."),
  ("Fuel quality & O&M plan","Fuel standard (e.g. ASTM D975 diesel), testing frequency and the generator test/maintenance regime."),
  ("Professional seals","Mechanical/structural seals for the tank foundation and containment.")],
 ["Secondary containment at 110% of the largest tank is non-negotiable — show the volume calculation.",
  "Emissions and boundary noise must match what the EIA approved.",
  "Above- vs. below-ground tanks change the rules; confirm the DOEB requirements for your choice."]),

("15. Telecommunications / Fiber Connection",
 [("Authority","National Broadcasting & Telecommunications Commission (NBTC) — if providing licensed service"),
  ("Phase","Construction"),("Lead time","~45 days"),("Est. cost","~5,000–15,000 THB"),
  ("Conditional","Only if you operate/resell telecom service; pure carrier connections may not need a licence")],
 "Covers either an NBTC licence (if the data centre provides telecom/ISP service) or, more commonly, the carrier connection "
 "arrangements bringing redundant fibre into the building.",
 [("Service application / licence (if applicable)","NBTC application for the relevant licence class if you provide public telecom service, "
   "with the company and financial-capacity evidence."),
  ("Network architecture & topology","Backbone design, redundancy, fibre routing into the building and entry points."),
  ("Technical specs & SLAs","Bandwidth, uptime targets and emergency procedures."),
  ("Security plan","Network and physical security for the telecom infrastructure."),
  ("Carrier agreements & routing","Contracts with fibre providers and the diverse routes/entries (for resilience).")],
 ["Most data centres only consume carrier service (no NBTC licence) — confirm whether you actually trigger licensing before filing.",
  "Design at least two diverse fibre entries for resilience; show the diversity.",
  "Align the entry points with the building and site plans."]),

("16. Electrical Inspection & Energisation",
 [("Authority","PEA / MEA"),
  ("Phase","Construction"),("Lead time","~14 days"),("Est. cost","~5,000 THB"),
  ("Depends on","Grid connection, substation, internal installation complete")],
 "The final inspection and authority-to-energise the electrical installation once construction is complete — the last step before the "
 "site has live utility power.",
 [("As-built electrical design","Final single-line, panel schedules, cable schedules and grounding diagram reflecting what was actually built."),
  ("Test reports","Grounding-resistance, insulation-resistance (megger), protection-coordination verification and, where used, transformer "
   "oil and load-bank tests."),
  ("Equipment certifications","Type/compliance certificates for switchgear, transformers and protective devices."),
  ("Compliance statement","Confirmation the installation meets the Thai Electrical Code (TIS), Building Code and utility rules, signed by a "
   "licensed electrical engineer."),
  ("Authority to energise","The qualified electrician's sign-off and emergency-shutdown/contact information.")],
 ["Have the test reports complete before booking the inspection — missing megger/earth results are the usual delay.",
  "As-builts must match the field installation exactly.",
  "Schedule against the commissioning programme so energisation doesn't hold up testing."]),

("17. Fire Safety Inspection",
 [("Authority","Local Fire Authority / Building Control"),
  ("Phase","Commissioning"),("Lead time","~7 days"),("Est. cost","No fee"),
  ("Depends on","Fire system installed; precedes occupancy")],
 "Physical verification that the installed fire systems work as designed — a precondition for the occupancy certificate.",
 [("System test certificates","Detection sensitivity tests, manual-call-point function, suppression pressure/discharge tests and agent-quantity verification."),
  ("Emergency lighting & signage","Illumination-level and sign-visibility checks."),
  ("Egress & passive protection","Clear escape routes, exit-sign illumination, and the integrity of fire-rated doors/walls/compartments."),
  ("Extinguishers & equipment","Quantity, location, charge and inspection tags."),
  ("Procedures & training","Posted evacuation procedures and staff fire-training records.")],
 ["Commission and test the suppression system before the inspection — re-tests cost schedule.",
  "Keep the maintenance/test log from day one; inspectors check it.",
  "Verify compartmentation penetrations are properly fire-stopped — a common failure."]),

("18. Building Completion Inspection",
 [("Authority","Local Building Control Authority"),
  ("Phase","Commissioning"),("Lead time","~7 days"),("Est. cost","~3,000 THB"),
  ("Precedes","Occupancy certificate")],
 "Confirms the building was constructed in accordance with the approved permit and is structurally and functionally complete.",
 [("As-built drawings","Architectural, structural and MEP as-builts with any deviations from the approved set noted and approved."),
  ("Professional completion certificates","Sign-off from the licensed architect and structural/MEP engineers that work matches the design."),
  ("Test & material certificates","Concrete cube/compression tests, steel/material certs and system commissioning reports."),
  ("Systems inspection","Structural elements, MEP, fire and electrical verified, with defect-rectification evidence."),
  ("Compliance checklist","Building-code compliance statement signed by the responsible engineer.")],
 ["Document any design changes during construction and get them approved before the inspection — undocumented deviations stall sign-off.",
  "Have commissioning reports ready; completion overlaps with system commissioning.",
  "Resolve the defect list before booking the inspection."]),

("19. Occupancy Certificate",
 [("Authority","Local Building Control Authority"),
  ("Phase","Commissioning"),("Lead time","~14 days"),("Est. cost","~2,000 THB"),
  ("Gateway","To lawful operation")],
 "The legal authorisation to occupy and use the building — the gateway from construction to operations. It bundles the prior inspections.",
 [("Completion, fire & electrical approvals","Copies of the building-completion, fire-safety and electrical-energisation sign-offs."),
  ("Occupancy application & as-builts","The use/occupancy application with final as-built drawings and the building-use statement."),
  ("Utility connection evidence","Proof electricity, water, wastewater and telecom are connected and operational."),
  ("Legal & insurance","Title/lease, tax registration and building insurance where required."),
  ("Operational procedures","Fire/emergency procedures, contacts and the building management plan.")],
 ["All upstream inspections must be passed first — occupancy is the assembly point, not a shortcut.",
  "Any change from the approved design must be documented and approved before this is granted.",
  "Confirm every utility is actually live, not merely contracted."]),

("20. Commissioning & Utility Acceptance",
 [("Authority","Utilities (PEA/MEA, MWA/PWA) and local authorities"),
  ("Phase","Commissioning"),("Lead time","~21 days"),("Est. cost","~15,000–30,000 THB"),
  ("Confirms","All systems operational and accepted")],
 "Final testing and formal acceptance that every utility and building system performs to specification before handover to operations.",
 [("System commissioning reports","Performance test results for electrical, water, wastewater, HVAC, fire and IT infrastructure, including "
   "backup-power (generator load-bank) tests."),
  ("Capacity verification","Load tests (electrical), flow tests (water), treatment-efficiency tests (wastewater) confirming rated capacity."),
  ("Utility acceptance sign-offs","The utilities' formal acceptance documents."),
  ("As-built system diagrams & O&M","Final diagrams, service manuals and maintenance schedules."),
  ("Operator readiness","Training records and 24/7 monitoring/support procedures for the operations team.")],
 ["Run a full integrated/black-building test (utility failure → generator → UPS) and record it.",
  "Acceptance is per-utility — track each sign-off separately.",
  "Hand over O&M documentation and training before go-live, not after."]),

("21. Factory License",
 [("Authority","Department of Industrial Works (DIW)"),
  ("Phase","Operations"),("Lead time","~30 days"),("Est. cost","~5,000 THB"),
  ("When needed","If the facility/machinery meets the Factory Act thresholds")],
 "Licenses the facility as a 'factory' under the Factory Act where its machinery/operations cross the statutory threshold — a data "
 "centre's plant (generators, cooling, electrical) often qualifies. Confirm applicability with DIW.",
 [("Factory licence application","The DIW application describing the operation, machinery list (capacity, power) and process."),
  ("Building/occupancy evidence","Building inspection/occupancy certificate showing the premises are lawful."),
  ("Environmental compliance","EIA approval or the relevant environmental notification, plus air/water/noise compliance."),
  ("Safety & worker plan","Fire safety, first aid, PPE, worker numbers/classifications and occupational-health procedures."),
  ("Waste & byproduct management","Types/quantities of waste and disposal methods, with any waste-disposal permits."),
  ("Manager appointment & insurance","Designated factory manager details and workers' compensation insurance.")],
 ["Confirm with DIW whether your plant actually triggers the Factory Act before assuming you need (or don't need) this.",
  "Machinery safety certificates and the electrical-safety inspection must be in hand.",
  "Waste/hazardous-material handling must align with the EIA."]),

("22. PDPA Compliance",
 [("Authority","Personal Data Protection Committee (PDPC)"),
  ("Phase","Operations"),("Lead time","~30 days to stand up"),("Est. cost","No fee"),
  ("Nature","Ongoing compliance, not a one-off licence")],
 "Establishes the data-protection framework required under Thailand's PDPA for any personal data the facility processes — an ongoing "
 "obligation rather than a single permit.",
 [("Data protection policy & privacy notices","Written policy and the notices given to data subjects describing processing purposes and rights."),
  ("Records of processing & data mapping","What personal data is held, where, why and for how long (retention schedule)."),
  ("Security measures","Technical and organisational controls — encryption, access control, logging, backup — protecting the data."),
  ("Data subject rights procedures","Processes to handle access, correction, deletion, portability and objection requests."),
  ("Processor/vendor agreements","Data-processing agreements with third parties (cloud, maintenance) including confidentiality clauses."),
  ("Breach response plan","Detection, remediation and the 72-hour PDPC notification procedure, with staff training records.")],
 ["Appoint a Data Protection Officer if your processing scale requires one.",
  "Breach notification to the PDPC is time-bound (within 72 hours) — rehearse it.",
  "As a colocation/host, clarify controller vs. processor roles with customers in contracts."]),

("23. Cybersecurity Compliance (CII)",
 [("Authority","National Cyber Security Agency (NCSA)"),
  ("Phase","Operations"),("Lead time","~30 days to stand up"),("Est. cost","No fee"),
  ("Nature","Ongoing; applies if designated Critical Information Infrastructure")],
 "Security governance required under the Cybersecurity Act, particularly if the facility is designated Critical Information "
 "Infrastructure (CII). Ongoing, with periodic reporting.",
 [("CII registration (if designated)","Notification/registration with NCSA where the facility is classed as CII."),
  ("Cybersecurity risk assessment","Asset inventory, threat and vulnerability assessment and impact analysis."),
  ("Policies & plans","Information-security policy, access-control policy, incident-response, business-continuity and disaster-recovery plans."),
  ("Technical controls","Network segmentation, firewalls/IDS, encryption, authentication, logging/monitoring (SIEM/SOC) and endpoint protection."),
  ("Testing & assurance","Regular penetration tests and vulnerability scans, and a recognised certification (ISO/IEC 27001, SOC 2)."),
  ("Reporting & supply-chain security","Incident-reporting procedures to NCSA, periodic compliance reporting, and hardware/software "
   "procurement and firmware-integrity controls.")],
 ["ISO/IEC 27001 also feeds the BOI tax-activation check — sequence the certification accordingly.",
  "Incident reporting to NCSA is mandatory and time-sensitive for CII — define the runbook now.",
  "Supply-chain/firmware integrity is increasingly scrutinised; document procurement controls."]),
]

for title, fct, purpose, items, pits in PERMITS:
    doc.add_page_break()
    H(title, 1)
    facts(fct)
    para("Purpose", bold=True, color=DARK, space_after=2)
    para(purpose, space_after=8)
    para("What each required item actually means", bold=True, color=DARK, space_after=2)
    for n,e in items:
        item(n,e)
    doc.add_paragraph()
    pitfalls(pits)

# ============================ APPENDIX ============================
doc.add_page_break()
H("Appendix · General review checklist before any submission", 1)
for x in [
 "Entity documents current (company affidavit & shareholder list issued within 3–6 months).",
 "Every foreign document apostilled/legalised AND translated to Thai by a certified translator.",
 "Certified true copies signed and sealed, not plain photocopies.",
 "POA references the exact application and is properly executed (notarised/apostilled if signed abroad).",
 "All engineering/architectural drawings sealed by Thai-licensed professionals with licence numbers.",
 "Technical figures (load MW, water demand, emissions, capital) consistent across BOI, EIA, utilities and building sets.",
 "Prerequisite permits attached where required (e.g. EIA letter with the building permit; ERC letter with BOI).",
 "Fresh copies pulled just before lodging to stay within validity windows.",
]:
    p = doc.add_paragraph(style='List Bullet'); p.add_run(x)

doc.add_paragraph()
para("Prepared from the CHN1A permitting research package (Master Permitting Guide, EIA & FBL detailed guides, "
     "BOI updates analysis, and the 23 permit checklists). Figures are planning estimates — confirm current "
     "requirements with each authority before lodging.", italic=True, color=GREY, size=9)

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "08_PERMIT_REQUIREMENTS_EXPLAINED.docx")
doc.save(out)
print("Saved:", out)
print("Permits documented:", len(PERMITS))
